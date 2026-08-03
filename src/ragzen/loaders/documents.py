"""Unified loader for common business document formats."""

from __future__ import annotations

import json
from html.parser import HTMLParser
from pathlib import Path

from ragzen.exceptions import MissingOptionalDependencyError, UnsupportedFileTypeError
from ragzen.loaders.base import compute_file_hash, validate_file
from ragzen.loaders.text import TextLoader
from ragzen.models import Document


class _VisibleTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._hidden_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        if tag in {"script", "style", "noscript"}:
            self._hidden_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._hidden_depth:
            self._hidden_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._hidden_depth and data.strip():
            self.parts.append(data.strip())


class UniversalDocumentLoader:
    """Load TXT, Markdown, CSV, JSON, HTML, PDF, DOCX and XLSX files."""

    _SUPPORTED_SUFFIXES = {
        ".txt",
        ".md",
        ".markdown",
        ".csv",
        ".json",
        ".html",
        ".htm",
        ".pdf",
        ".docx",
        ".xlsx",
    }

    def __init__(self, *, max_size_mb: float = 100.0) -> None:
        self._max_size_mb = max_size_mb
        self._text_loader = TextLoader(max_size_mb=max_size_mb)

    def supported_mime_types(self) -> frozenset[str]:
        return frozenset(
            {
                "text/plain",
                "text/markdown",
                "text/csv",
                "application/json",
                "text/html",
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }
        )

    def load(self, source: str | Path) -> list[Document]:
        path = Path(source)
        suffix = path.suffix.casefold()
        if suffix not in self._SUPPORTED_SUFFIXES:
            raise UnsupportedFileTypeError(f"Unsupported document type: {suffix or '<none>'}")
        validate_file(path, max_size_mb=self._max_size_mb)
        if suffix in {".txt", ".md", ".markdown", ".csv"}:
            return self._text_loader.load(path)
        content, page_count = self._extract(path, suffix)
        return [
            Document(
                tenant_id="__unassigned__",
                content=content,
                content_hash=compute_file_hash(path),
                source="file",
                source_uri=str(path.resolve()),
                file_name=path.name,
                mime_type=self._mime_type(suffix),
                page_count=page_count,
                document_type=suffix.lstrip("."),
            )
        ]

    def _extract(self, path: Path, suffix: str) -> tuple[str, int]:
        if suffix == ".json":
            data = json.loads(path.read_text(encoding="utf-8"))
            return json.dumps(data, ensure_ascii=False, indent=2), 1
        if suffix in {".html", ".htm"}:
            parser = _VisibleTextParser()
            parser.feed(path.read_text(encoding="utf-8", errors="replace"))
            return "\n".join(parser.parts), 1
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:
                raise MissingOptionalDependencyError("pypdf", "documents", "PDF loading") from exc
            reader = PdfReader(str(path))
            content = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            return content, len(reader.pages)
        if suffix == ".docx":
            try:
                from docx import Document as DocxDocument
            except ImportError as exc:
                raise MissingOptionalDependencyError(
                    "python-docx", "documents", "DOCX loading"
                ) from exc
            document = DocxDocument(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs), 1
        if suffix == ".xlsx":
            try:
                from openpyxl import load_workbook
            except ImportError as exc:
                raise MissingOptionalDependencyError(
                    "openpyxl", "documents", "XLSX loading"
                ) from exc
            workbook = load_workbook(path, read_only=True, data_only=True)
            parts: list[str] = []
            try:
                for worksheet in workbook.worksheets:
                    parts.append(f"# Sheet: {worksheet.title}")
                    for row in worksheet.iter_rows(values_only=True):
                        values = [str(value) for value in row if value is not None]
                        if values:
                            parts.append("\t".join(values))
            finally:
                workbook.close()
            return "\n".join(parts), len(workbook.sheetnames)
        raise UnsupportedFileTypeError(f"Unsupported document type: {suffix}")

    @staticmethod
    def _mime_type(suffix: str) -> str:
        return {
            ".json": "application/json",
            ".html": "text/html",
            ".htm": "text/html",
            ".pdf": "application/pdf",
            ".docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }[suffix]
